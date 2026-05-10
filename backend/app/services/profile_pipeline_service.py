"""Background pipeline for building the user's company profile.

The company profile is used in two different contexts:
- Frontend display: concise Chinese fields for the user to inspect quickly.
- Outreach generation: complete source-backed data, including an English copy.
"""

from __future__ import annotations

import asyncio
import base64
import csv
import importlib.util
import io
import json
import logging
import os
import re
import time
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

import httpx
import replicate
from sqlalchemy import select, update

from app.config import settings
from app.database import async_session_factory
from app.models.company_profile import CompanyProfile
from app.models.task import Task, TaskLog
from app.services import task_manager
from app.services.settings_service import ensure_recommended_email_settings
from app.utils.db_sequences import sync_company_profiles_id_sequence

logger = logging.getLogger(__name__)

if settings.replicate_api_token and not os.environ.get("REPLICATE_API_TOKEN"):
    os.environ["REPLICATE_API_TOKEN"] = settings.replicate_api_token


PROJECT_ROOT = Path(__file__).resolve().parents[3]
SCRAPER_PATH = (
    PROJECT_ROOT
    / "waimao_toolkit_new"
    / "skills"
    / "company-profile"
    / "scripts"
    / "scrape_website.py"
)


PROFILE_SYSTEM_PROMPT = """你是一位资深 B2B 销售资料顾问，负责按照 company-profile skill 构建企业销售能力档案。

目标：把用户提供的信息和官网抓取内容整理成一份后续客户开发、客户匹配、开发信生成都可以直接调用的结构化公司画像。

输出语言规则：
1. 根级主字段必须用中文，给前端展示使用。
2. 同时生成 english_profile，保留英文版，用于后续英文开发信。
3. 专有名词、认证名称、客户/项目英文原名可以保留英文。

真实性规则：
1. 不只写公司简介，要构建"销售认知底座"：我们是谁、卖什么、能帮客户解决什么、适合开发什么客户。
2. 信息必须具体、可验证。不要写"质量好、服务好"这类空话，除非能转成具体证据。
3. 不确定的信息不要编造；缺失字段留空、空数组，或在 metadata.notes 标注"待确认"。
4. 案例、认证、产能、客户名称、交付数据不能凭空生成。
5. case_studies 只有在资料中找到明确项目/案例线索时才填写；没有就返回空数组。
6. boundaries 必须明确写出可以说、不能乱说、敏感话题。
7. 输出必须是严格 JSON，不要 Markdown 代码块，不要额外解释。

画像质量标准：
1. 不只总结官网首页，要把资料转成销售可用的认知底座：产品适合谁、解决什么问题、为什么可信、开发信里怎么说。
2. 必须尽量补齐 skill 文档要求的关键维度：基础信息、主营产品、核心竞争力、适合开发的客户类型、成功案例、证书资质、合作模式、独特卖点、客户匹配建议、信息边界、metadata。
3. case_studies 必须尽量从 Projects / Cases / Gallery / Portfolio / Applications / Solutions / Clients 等资料中深挖；资料充分时目标至少 10 个案例，同类行业/地区只保留最有代表性的 2-3 个；资料不足时不要编造，在 metadata.notes 说明待补充。
4. 每个案例优先补齐 project、project_en、client_type、industry、country、products_used、area_or_quantity、problem_solved、result、key_highlight、usable_in_outreach。
5. key_highlight 要能直接放进开发信，优先包含数字、地区、产品、交付结果；没有明确数据时不要编造，写可验证的简短亮点。
6. customer_matching_guide 必须站在客户开发视角，说明哪类客户最值得开发、开发信优先强调什么、哪些话题要避免。
7. boundaries 必须防止后续开发信乱写，明确 claims_we_can_make、claims_we_cannot_make、sensitive_topics。
8. metadata 必须记录 source_urls、source_documents、profile_completeness 和仍需确认的 notes。
9. profile_completeness 必须按已填字段真实评估：0.0-0.3=基础信息不足；0.3-0.6=有基本框架；0.6-0.8=较完整可用于开发信；0.8-1.0=覆盖关键维度。只要 products、case_studies、core_competencies 等已有实质内容，不得返回 0。

返回 JSON 结构：
{
  "company_name": "",
  "one_line_intro": "",
  "full_intro": "",
  "location": "",
  "industry": "",
  "established": "",
  "scale": "",
  "website": "",
  "products": [
    {
      "name": "",
      "description": "",
      "target_customers": "",
      "key_selling_points": []
    }
  ],
  "core_competencies": [
    {
      "competency": "",
      "description": "",
      "evidence": ""
    }
  ],
  "target_customer_types": [
    {
      "type": "",
      "why_suitable": "",
      "pitch_focus": []
    }
  ],
  "case_studies": [
    {
      "project": "",
      "project_en": "",
      "client_type": "",
      "industry": "",
      "country": "",
      "products_used": [],
      "area_or_quantity": "",
      "problem_solved": "",
      "result": "",
      "key_highlight": "",
      "usable_in_outreach": false
    }
  ],
  "certifications": [],
  "cooperation_models": [
    {
      "model": "",
      "description": "",
      "customer_value": ""
    }
  ],
  "unique_selling_points": [],
  "customer_matching_guide": [
    {
      "customer_type": "",
      "priority_points": [],
      "avoid_topics": []
    }
  ],
  "boundaries": {
    "claims_we_can_make": [],
    "claims_we_cannot_make": [],
    "sensitive_topics": []
  },
  "english_profile": {
    "company_name": "",
    "one_line_intro": "",
    "full_intro": "",
    "products": [],
    "core_competencies": [],
    "case_studies": [],
    "unique_selling_points": [],
    "customer_matching_guide": []
  },
  "metadata": {
    "source_urls": [],
    "source_documents": [],
    "profile_completeness": 0.0,
    "notes": ""
  }
}"""

PROFILE_EDIT_SYSTEM_PROMPT = """你是企业资料编辑助手。用户会给你一份已有的公司画像 JSON 和修改指令，你需要严格按 company-profile skill 的更新模式修改画像。

核心规则：
1. 默认基于现有画像做增量补充/修正，不要从零重写，不要覆盖未涉及内容。
2. 用户要求删除/移除/不要/去掉/剔除 → 从返回 JSON 中彻底移除对应条目，不要保留。
3. 用户要求添加/补充 → 按 company-profile schema 在对应位置添加新条目，并尽量补齐销售可用字段。
4. 用户要求修改/更正/改为 → 按指令修改。
5. 未被指令涉及的部分保持完全不变，一个字段都不要自行修改。
6. 如果修改 products、case_studies、core_competencies、target_customer_types、cooperation_models、customer_matching_guide、boundaries，必须按 skill 文档的完整度要求补齐字段；信息不足时留空或在 metadata.notes 标注待确认，不要编造。
7. 如果画像很长，你只需要返回被修改的字段/数组，不需要返回完整画像。
   - 例如只修改了 certifications，就只返回 {"certifications": [...修改后的数组...]}
   - 修改了多个字段就都返回，如 {"certifications": [...], "one_line_intro": "..."}
8. 同步更新 metadata.updated_at；如修改影响完整度，更新 metadata.profile_completeness。
9. 前端展示字段用中文，english_profile 用英文；修改英文开发信会用到的字段时同步更新 english_profile 对应部分。
10. 不编造用户没有提供的信息。
11. 直接返回 JSON，不要 Markdown 代码块"""


async def _create_task_log(
    db,
    task_id: int,
    step: int,
    name: str,
    status: str = "running",
    message: str = "",
    progress: int = 0,
) -> TaskLog:
    log = TaskLog(
        task_id=task_id,
        step_number=step,
        step_name=name,
        status=status,
        message=message,
        progress=progress,
    )
    db.add(log)
    await db.commit()
    await db.refresh(log)
    return log


async def _update_task_log(
    db,
    task_id: int,
    step: int,
    status: str | None = None,
    message: str | None = None,
    progress: int | None = None,
) -> None:
    result = await db.execute(
        select(TaskLog).where(
            TaskLog.task_id == task_id,
            TaskLog.step_number == step,
        )
    )
    log = result.scalar_one_or_none()
    if log is None:
        return
    if status is not None:
        log.status = status
    if message is not None:
        log.message = message
    if progress is not None:
        log.progress = progress
    await db.commit()


def extract_url(text: str) -> str:
    """Extract the first website URL from a user message."""
    if not text:
        return ""
    match = re.search(r"https?://[A-Za-z0-9\-._~:/?#\[\]@!$&'()*+,;=%]+", text)
    if match:
        return match.group(0).rstrip(").,;，。；")
    domain_match = re.search(
        r"\b(?:www\.)?[a-zA-Z0-9-]+(?:\.[a-zA-Z0-9-]+)+\b",
        text,
    )
    if domain_match:
        domain = domain_match.group(0).rstrip(").,;，。；")
        return domain if domain.startswith("http") else f"https://{domain}"
    return ""


def _normalize_url(url: str) -> str:
    if not url:
        return ""
    url = url.strip()
    if not urlparse(url).scheme:
        url = f"https://{url}"
    return url.rstrip("/")


def _load_scraper_module():
    spec = importlib.util.spec_from_file_location("company_profile_scraper", SCRAPER_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load scraper module: {SCRAPER_PATH}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _scrape_once(scraper, url: str, max_pages: int) -> list[dict]:
    if scraper.HAS_PLAYWRIGHT:
        return scraper.scrape_website_playwright(url, max_pages=max_pages)
    return scraper.scrape_website_requests(url, max_pages=max_pages)


def _scrape_candidates(url: str) -> list[str]:
    normalized_url = _normalize_url(url)
    candidates = [normalized_url]
    parsed = urlparse(normalized_url)
    if parsed.scheme in {"http", "https"}:
        alternate_scheme = "https" if parsed.scheme == "http" else "http"
        alternate = parsed._replace(scheme=alternate_scheme).geturl()
        if alternate not in candidates:
            candidates.append(alternate)
    return candidates


def _scrape_website(url: str, max_pages: int = 6) -> tuple[str, list[str]]:
    """Scrape a website using the company-profile skill scraper functions."""
    scraper = _load_scraper_module()
    mode = "playwright" if scraper.HAS_PLAYWRIGHT else "requests"
    last_error: Exception | None = None

    for candidate in _scrape_candidates(url):
        try:
            pages_data = _scrape_once(scraper, candidate, max_pages)
            if not pages_data:
                raise RuntimeError(f"scraper returned 0 pages for {candidate}")
            markdown = scraper.format_as_markdown(pages_data)
            urls = [p.get("url", "") for p in pages_data if p.get("url")]
            logger.info(
                "Company profile scrape succeeded: mode=%s url=%s pages=%d",
                mode,
                candidate,
                len(urls),
            )
            return markdown, urls
        except Exception as exc:
            last_error = exc
            logger.warning(
                "Company profile scrape attempt failed: mode=%s url=%s error=%s",
                mode,
                candidate,
                exc,
            )

    raise RuntimeError(f"官网没有成功抓取任何页面：{last_error}")


def _compact_scraped_markdown(markdown: str, max_chars: int = 12000) -> str:
    """Keep AI input focused on useful English/source pages.

    Some sites expose many translated copies of the same page. Sending those
    copies to the model makes the profile slower without adding much signal.
    """
    if not markdown:
        return ""

    language_path_pattern = re.compile(r"Source:\s+\S+/(?:vi|ur|tr|tl|ar|fr|es|de|it|pt|ru|ja|ko)/", re.I)
    sections = markdown.split("\n---\n")
    kept: list[str] = []
    for section in sections:
        if language_path_pattern.search(section):
            continue
        kept.append(section)

    compacted = "\n---\n".join(kept).strip() or markdown
    return compacted[:max_chars]


async def _run_blocking_step_with_progress(
    *,
    db,
    task_id: int,
    step: int,
    func,
    args: tuple = (),
    timeout_seconds: int,
    running_message: str,
    timeout_message: str,
):
    """Run blocking work while keeping SSE task logs fresh.

    The frontend marks a task stale when no task_log update is emitted for too
    long. Scraping and model calls are blocking, so this wrapper advances the
    current step while the worker thread is still running.
    """
    worker = asyncio.create_task(asyncio.to_thread(func, *args))
    started = time.monotonic()
    progress = 10

    while True:
        try:
            return await asyncio.wait_for(asyncio.shield(worker), timeout=10)
        except asyncio.TimeoutError:
            elapsed = time.monotonic() - started
            if elapsed >= timeout_seconds:
                worker.cancel()
                raise TimeoutError(timeout_message)

            progress = min(90, progress + 10)
            await _update_task_log(
                db,
                task_id,
                step,
                message=f"{running_message}（已运行 {int(elapsed)} 秒）",
                progress=progress,
            )
            task_manager.update_heartbeat(task_id)


def _parse_ai_json(response_text: str) -> dict | None:
    text = response_text.strip()
    match = re.search(r"```(?:json)?\s*\n?([\s\S]*?)\n?```", text)
    if match:
        text = match.group(1).strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        try:
            return json.loads(text[start : end + 1])
        except json.JSONDecodeError:
            return None
    return None


def _ensure_list(value) -> list:
    if isinstance(value, list):
        return value
    if value in (None, ""):
        return []
    return [value]


def _normalize_completeness(value) -> float:
    try:
        score = float(value or 0)
    except (TypeError, ValueError):
        return 0.0
    if score > 1:
        score = score / 100
    return max(0.0, min(score, 1.0))


def _has_text(value) -> bool:
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, (int, float, bool)):
        return True
    if isinstance(value, dict):
        return any(_has_text(v) for v in value.values())
    if isinstance(value, list):
        return any(_has_text(v) for v in value)
    return False


def _nonempty_count(value) -> int:
    return sum(1 for item in _ensure_list(value) if _has_text(item))


def _derive_profile_completeness(profile: dict) -> float:
    """Derive a conservative completeness score from populated profile fields."""
    if not isinstance(profile, dict):
        return 0.0

    score = 0.0

    basic_fields = [
        "company_name",
        "one_line_intro",
        "full_intro",
        "industry",
        "location",
        "website",
        "scale",
        "established",
    ]
    basic_filled = sum(1 for key in basic_fields if _has_text(profile.get(key)))
    score += min(0.2, basic_filled * 0.025)

    products = _nonempty_count(profile.get("products"))
    competencies = _nonempty_count(profile.get("core_competencies"))
    target_types = _nonempty_count(profile.get("target_customer_types"))
    cases = _nonempty_count(profile.get("case_studies"))
    certs = _nonempty_count(profile.get("certifications"))
    models = _nonempty_count(profile.get("cooperation_models"))
    usps = _nonempty_count(profile.get("unique_selling_points"))
    guide = _nonempty_count(profile.get("customer_matching_guide"))

    score += min(0.18, products * 0.045)
    score += min(0.14, competencies * 0.04 + usps * 0.02)
    score += min(0.12, target_types * 0.04)
    score += min(0.18, cases * 0.06)
    score += min(0.1, guide * 0.04)
    score += min(0.08, certs * 0.025 + models * 0.03)

    boundaries = profile.get("boundaries") if isinstance(profile.get("boundaries"), dict) else {}
    boundary_filled = sum(1 for key in ("claims_we_can_make", "claims_we_cannot_make", "sensitive_topics") if _has_text(boundaries.get(key)))
    score += min(0.05, boundary_filled * 0.018)

    english = profile.get("english_profile") if isinstance(profile.get("english_profile"), dict) else {}
    if _has_text(english.get("one_line_intro")) or _has_text(english.get("products")):
        score += 0.05

    metadata = profile.get("metadata") if isinstance(profile.get("metadata"), dict) else {}
    if _has_text(metadata.get("source_urls")):
        score += 0.05

    return round(max(0.0, min(score, 0.95)), 2)


def _compact_join(items: list) -> str:
    values = []
    for item in items:
        if isinstance(item, str):
            values.append(item)
        elif isinstance(item, dict):
            values.append(
                item.get("name")
                or item.get("model")
                or item.get("competency")
                or item.get("customer_type")
                or item.get("type")
                or ""
            )
    return "、".join(v for v in values if v)


def _normalize_profile(profile: dict, website: str, source_urls: list[str]) -> dict:
    """Keep the profile close to the skill schema and fill safe defaults."""
    now = datetime.now().isoformat(timespec="seconds")
    profile = profile if isinstance(profile, dict) else {}

    for key in (
        "products",
        "core_competencies",
        "target_customer_types",
        "case_studies",
        "certifications",
        "cooperation_models",
        "unique_selling_points",
        "customer_matching_guide",
    ):
        profile[key] = _ensure_list(profile.get(key))

    boundaries = profile.get("boundaries")
    if not isinstance(boundaries, dict):
        boundaries = {}
    boundaries["claims_we_can_make"] = _ensure_list(boundaries.get("claims_we_can_make"))
    boundaries["claims_we_cannot_make"] = _ensure_list(boundaries.get("claims_we_cannot_make"))
    boundaries["sensitive_topics"] = _ensure_list(boundaries.get("sensitive_topics"))
    profile["boundaries"] = boundaries

    english_profile = profile.get("english_profile")
    if not isinstance(english_profile, dict):
        english_profile = {}
    english_profile.setdefault("company_name", profile.get("company_name", ""))
    english_profile.setdefault("one_line_intro", profile.get("one_line_intro", ""))
    english_profile.setdefault("full_intro", profile.get("full_intro", ""))
    for key in (
        "products",
        "core_competencies",
        "case_studies",
        "unique_selling_points",
        "customer_matching_guide",
    ):
        english_profile[key] = _ensure_list(english_profile.get(key))
    profile["english_profile"] = english_profile

    metadata = profile.get("metadata")
    if not isinstance(metadata, dict):
        metadata = {}
    metadata.setdefault("created_at", now)
    metadata["updated_at"] = now
    metadata["source_urls"] = list(dict.fromkeys(_ensure_list(metadata.get("source_urls")) + source_urls))
    metadata.setdefault("source_documents", [])
    ai_completeness = _normalize_completeness(
        metadata.get("profile_completeness", 0.0)
    )
    metadata["profile_completeness"] = max(ai_completeness, _derive_profile_completeness(profile))
    metadata.setdefault("notes", "")
    profile["metadata"] = metadata

    if website and not profile.get("website"):
        profile["website"] = website
    profile.setdefault("company_name", "")
    profile.setdefault("one_line_intro", "")
    profile.setdefault("full_intro", "")
    profile.setdefault("location", "")
    profile.setdefault("industry", "")
    profile.setdefault("established", "")
    profile.setdefault("scale", "")
    profile.setdefault("website", website)

    return profile


def _is_replace_profile_request(text: str) -> bool:
    text = text or ""
    replace_terms = [
        "重新采集",
        "重新生成",
        "重新整理",
        "从头",
        "覆盖",
        "不要之前",
        "不要原来",
        "全部重做",
    ]
    supplement_terms = [
        "补充",
        "补一下",
        "补充一下",
        "修改",
        "修正",
        "更新",
        "加上",
        "追加",
        "这些也是",
        "这是案例",
        "这些是案例",
    ]
    if any(term in text for term in supplement_terms):
        return False
    return any(term in text for term in replace_terms)


def _dedupe_key(value: Any) -> str:
    if isinstance(value, dict):
        for key in ("name", "model", "competency", "customer_type", "type", "project", "title"):
            if value.get(key):
                return f"{key}:{str(value[key]).strip().lower()}"
        return json.dumps(value, ensure_ascii=False, sort_keys=True).lower()
    return str(value).strip().lower()


def _merge_lists(old_value: Any, new_value: Any) -> list[Any]:
    merged: list[Any] = []
    seen: set[str] = set()
    for item in _ensure_list(old_value) + _ensure_list(new_value):
        key = _dedupe_key(item)
        if not key or key in seen:
            continue
        seen.add(key)
        merged.append(item)
    return merged


def _merge_profile_data(existing: dict, generated: dict) -> dict:
    """Merge generated supplemental profile data into the current profile."""
    if not isinstance(existing, dict):
        existing = {}
    if not isinstance(generated, dict):
        generated = {}

    merged = dict(existing)
    scalar_keys = [
        "company_name",
        "one_line_intro",
        "full_intro",
        "location",
        "industry",
        "established",
        "scale",
        "website",
    ]
    for key in scalar_keys:
        if generated.get(key):
            merged[key] = generated[key]
        else:
            merged.setdefault(key, existing.get(key, ""))

    for key in (
        "products",
        "core_competencies",
        "target_customer_types",
        "case_studies",
        "certifications",
        "cooperation_models",
        "unique_selling_points",
        "customer_matching_guide",
    ):
        merged[key] = _merge_lists(existing.get(key), generated.get(key))

    for key in ("boundaries", "english_profile", "metadata"):
        old_obj = existing.get(key) if isinstance(existing.get(key), dict) else {}
        new_obj = generated.get(key) if isinstance(generated.get(key), dict) else {}
        merged_obj = dict(old_obj)
        for sub_key, new_value in new_obj.items():
            old_value = old_obj.get(sub_key)
            if isinstance(old_value, list) or isinstance(new_value, list):
                merged_obj[sub_key] = _merge_lists(old_value, new_value)
            elif new_value not in (None, "", {}, []):
                merged_obj[sub_key] = new_value
            else:
                merged_obj.setdefault(sub_key, old_value)
        merged[key] = merged_obj

    notes = []
    old_notes = (existing.get("metadata") or {}).get("notes", "") if isinstance(existing.get("metadata"), dict) else ""
    new_notes = (generated.get("metadata") or {}).get("notes", "") if isinstance(generated.get("metadata"), dict) else ""
    if old_notes:
        notes.append(old_notes)
    if new_notes and new_notes not in notes:
        notes.append(new_notes)
    merged.setdefault("metadata", {})
    if isinstance(merged["metadata"], dict):
        merged["metadata"]["notes"] = "\n".join(notes)

    return merged


def _build_timeout_fallback_profile(
    params: dict,
    source_text: str,
    scraped_markdown: str,
    source_urls: list[str],
) -> dict:
    website = params.get("url", "")
    hostname = urlparse(website).netloc.replace("www.", "") if website else ""
    company_name = hostname.split(".")[0].upper() if hostname else "未命名公司"

    title_match = re.search(r"## \[[^\]]+\]\s+(.+)", scraped_markdown or "")
    if title_match:
        title = title_match.group(1).strip()
        if title and title.lower() != "untitled":
            company_name = title.split("|")[0].split("-")[0].strip()[:80] or company_name

    compact_text = _compact_scraped_markdown(scraped_markdown, max_chars=3000)
    compact_text = re.sub(r"#+\s*|\*\*|Source:\s*\S+", "", compact_text)
    compact_text = re.sub(r"\s+", " ", compact_text).strip()
    intro = compact_text[:700] if compact_text else source_text[:700]

    return _normalize_profile(
        {
            "company_name": company_name,
            "one_line_intro": intro[:220],
            "full_intro": intro,
            "website": website,
            "products": [],
            "core_competencies": [],
            "target_customer_types": [],
            "case_studies": [],
            "certifications": [],
            "cooperation_models": [],
            "unique_selling_points": [],
            "customer_matching_guide": [],
            "boundaries": {
                "claims_we_can_make": [
                    "仅可使用已抓取官网资料和用户已提供的信息进行表述。",
                ],
                "claims_we_cannot_make": [
                    "AI 整理阶段超时，暂不能确认资质、案例、交付周期或性能参数。",
                ],
                "sensitive_topics": [
                    "客户案例、认证证书、工程性能和交付承诺需人工确认。",
                ],
            },
            "metadata": {
                "source_urls": source_urls,
                "source_documents": [],
                "profile_completeness": 0.25 if intro else 0.1,
                "notes": "AI 整理画像超时，系统已保存基础画像草稿；建议补充公司简介、主打产品、认证和案例核心文字后重新采集。",
            },
        },
        website,
        source_urls,
    )


def _profile_to_markdown(profile: dict) -> str:
    metadata = profile.get("metadata", {})
    lines = [
        f"# {profile.get('company_name') or '公司画像'} - 企业能力档案",
        "",
        f"> 更新时间：{metadata.get('updated_at', '')}",
        "",
        "## 公司概况",
        "",
        profile.get("one_line_intro", ""),
        "",
        profile.get("full_intro", ""),
        "",
        f"- 行业：{profile.get('industry', '')}",
        f"- 地区：{profile.get('location', '')}",
        f"- 官网：{profile.get('website', '')}",
        f"- 成立时间：{profile.get('established', '')}",
        f"- 规模：{profile.get('scale', '')}",
        "",
        "## 主营产品与服务",
        "",
    ]

    for product in profile.get("products", []):
        if not isinstance(product, dict):
            lines.extend([f"### {product}", ""])
            continue
        lines.extend(
            [
                f"### {product.get('name', '')}",
                f"- 描述：{product.get('description', '')}",
                f"- 适合客户：{product.get('target_customers', '')}",
                f"- 关键卖点：{_compact_join(product.get('key_selling_points', []))}",
                "",
            ]
        )

    lines.extend(["## 核心竞争力", ""])
    for item in profile.get("core_competencies", []):
        if isinstance(item, dict):
            lines.append(
                f"- {item.get('competency', '')}：{item.get('description', '')}（证据：{item.get('evidence', '')}）"
            )
        else:
            lines.append(f"- {item}")

    lines.extend(["", "## 适合开发的客户类型", ""])
    for item in profile.get("target_customer_types", []):
        if isinstance(item, dict):
            lines.extend(
                [
                    f"### {item.get('type', '')}",
                    f"- 为什么适合：{item.get('why_suitable', '')}",
                    f"- 开发信重点：{_compact_join(item.get('pitch_focus', []))}",
                    "",
                ]
            )

    lines.extend(["## 成功案例", ""])
    for case in profile.get("case_studies", []):
        if isinstance(case, dict):
            lines.extend(
                [
                    f"### {case.get('project', '')}",
                    f"- 英文项目名：{case.get('project_en', '')}",
                    f"- 客户类型：{case.get('client_type', '')}",
                    f"- 行业/地区：{case.get('industry', '')} / {case.get('country', '')}",
                    f"- 使用产品：{_compact_join(case.get('products_used', []))}",
                    f"- 规模：{case.get('area_or_quantity', '')}",
                    f"- 解决问题：{case.get('problem_solved', '')}",
                    f"- 结果：{case.get('result', '')}",
                    f"- 一句话亮点：{case.get('key_highlight', '')}",
                    f"- 可用于开发信：{'是' if case.get('usable_in_outreach') else '否'}",
                    "",
                ]
            )
    if not profile.get("case_studies"):
        lines.append("暂无明确可用案例。")

    lines.extend(["", "## 独特卖点", ""])
    for item in profile.get("unique_selling_points", []):
        lines.append(f"- {item if isinstance(item, str) else json.dumps(item, ensure_ascii=False)}")

    lines.extend(["", "## 合作模式", ""])
    for item in profile.get("cooperation_models", []):
        if isinstance(item, dict):
            lines.append(f"- {item.get('model', '')}：{item.get('description', '')}；客户价值：{item.get('customer_value', '')}")
        else:
            lines.append(f"- {item}")

    lines.extend(["", "## 资质认证", ""])
    for item in profile.get("certifications", []):
        lines.append(f"- {item if isinstance(item, str) else json.dumps(item, ensure_ascii=False)}")

    lines.extend(["", "## 信息边界", ""])
    boundaries = profile.get("boundaries", {})
    lines.append(f"- 可以说：{_compact_join(boundaries.get('claims_we_can_make', []))}")
    lines.append(f"- 不能乱说：{_compact_join(boundaries.get('claims_we_cannot_make', []))}")
    lines.append(f"- 敏感话题：{_compact_join(boundaries.get('sensitive_topics', []))}")

    if profile.get("english_profile"):
        lines.extend(["", "## English Version", ""])
        english = profile["english_profile"]
        lines.append(english.get("one_line_intro", ""))
        lines.append("")
        lines.append(english.get("full_intro", ""))

    return "\n".join(lines).strip() + "\n"


# Keywords → profile section mapping for smart section extraction
PROFILE_SECTION_KEYWORDS: dict[str, list[str]] = {
    "certifications": ["资质", "检测", "认证", "SGS", "ISO", "CE", "RoHS", "证书", "报告", "环保", "测试"],
    "case_studies": ["案例", "项目", "工程", "案例研究", "合作案例", "成功案例", "portfolio", "projects"],
    "products": ["产品", "商品", "货品", "品类", "型号"],
    "core_competencies": ["优势", "竞争力", "核心", "强项", "能力"],
    "target_customer_types": ["客户类型", "目标客户", "买家类型", "适合客户"],
    "cooperation_models": ["合作模式", "合作方式", "经销", "代理", "OEM", "ODM"],
    "unique_selling_points": ["卖点", "独特", "差异化", "亮点"],
    "customer_matching_guide": ["客户匹配", "匹配指南", "开发重点", "避免话题"],
    "boundaries": ["边界", "可以说", "不能说", "敏感", "禁止"],
    "english_profile": ["英文", "english"],
}

# Scalar fields that need special handling
PROFILE_SCALAR_FIELDS = [
    "company_name", "one_line_intro", "full_intro",
    "location", "industry", "established", "scale", "website",
]


def _identify_edit_sections(message: str) -> list[str]:
    """Identify which profile sections the user wants to modify based on keywords."""
    text = (message or "").lower()
    matched = set()
    for section, keywords in PROFILE_SECTION_KEYWORDS.items():
        if any(kw in text for kw in keywords):
            matched.add(section)
    # Also check scalar field names
    for field in PROFILE_SCALAR_FIELDS:
        if field.replace("_", "") in text.replace(" ", "") or field in text:
            matched.add(field)
    return sorted(matched)


def _profile_section_summary(profile: dict) -> str:
    """Build a compact summary of all profile sections for context."""
    lines = []
    for key in PROFILE_SCALAR_FIELDS:
        val = profile.get(key, "")
        if val:
            lines.append(f"  {key}: {val[:100]}")
    for section in PROFILE_SECTION_KEYWORDS:
        val = profile.get(section)
        if isinstance(val, list):
            lines.append(f"  {section}: [{len(val)} items]")
        elif isinstance(val, dict) and val:
            keys = list(val.keys())[:5]
            lines.append(f"  {section}: dict with keys {keys}")
    return "\n".join(lines)


def _build_user_prompt(params: dict, source_text: str, scraped_markdown: str, image_descriptions: str = "") -> str:
    website = params.get("url", "")
    user_message = params.get("source_text", "")
    existing_profile = params.get("existing_profile")

    # Edit mode: existing profile present → smart section extraction
    if isinstance(existing_profile, dict) and existing_profile:
        sections = _identify_edit_sections(user_message or source_text or "")

        parts = []
        parts.append("## 用户的修改指令")
        parts.append(user_message or source_text or "未提供")
        parts.append("")

        if sections:
            # Only send the relevant sections + a compact summary of everything else
            parts.append("## 画像结构概览（仅供参考，不要修改这些部分）")
            parts.append(_profile_section_summary(existing_profile))
            parts.append("")
            parts.append(f"## 需要修改的部分（只修改以下字段）")
            for section in sections:
                val = existing_profile.get(section)
                if val is not None:
                    section_json = json.dumps(val, ensure_ascii=False)
                    parts.append(f"### {section}")
                    parts.append(section_json)
        else:
            # Can't identify sections — send compact summary + ask AI to specify
            parts.append("## 画像结构概览")
            parts.append(_profile_section_summary(existing_profile))
            parts.append("")
            parts.append("(请根据用户指令判断需要修改哪个部分，返回修改后的完整 JSON)")

        if image_descriptions:
            parts.append("")
            parts.append("## 用户上传图片 AI 识别内容")
            parts.append(image_descriptions)
        if scraped_markdown:
            parts.append("")
            parts.append("## 新的网页资料")
            parts.append(_compact_scraped_markdown(scraped_markdown))
        parts.append("")
        parts.append("请按 company-profile skill 更新模式处理：只处理变化部分，保留未变化内容；新增案例/产品/客户类型/合作模式/边界字段时按完整 schema 补齐，信息不足写待确认，不要编造。")
        parts.append("请返回修改后的字段，只返回被修改的字段。格式如：{\"certifications\": [...], \"one_line_intro\": \"...\"}")
        parts.append("不要返回未修改的字段，不要返回完整画像，不要 Markdown 代码块。")
        return "\n".join(parts)

    # Create mode: original prompt structure
    image_block = ""
    if image_descriptions:
        image_block = f"\n=== 用户上传图片 AI 识别内容 ===\n{image_descriptions}\n"
    return f"""请根据以下资料，按照 company-profile skill 的标准生成企业能力档案。

特别注意：
- 前端展示字段请写中文，简洁自然。
- english_profile 请写英文，给后续开发信使用。
- 必须按 company-profile skill 的完整度标准输出：基础信息、主营产品、核心竞争力、适合开发的客户类型、案例、证书资质、合作模式、独特卖点、客户匹配建议、信息边界、metadata。
- 如果抓取资料中包含 Projects、Cases、Case Studies、Gallery、Portfolio、Applications、Solutions、Clients 等页面，请优先挖掘项目案例。
- 案例必须来自资料里的明确线索，不能编造客户名、国家、数量、认证或效果。
- 请严格按 company-profile skill 的销售资料标准处理：不要只写公司简介，要输出后续客户开发可直接复用的产品、客户类型、核心优势、案例亮点、合作模式、信息边界。
- case_studies 资料充分时目标至少 10 个；每条尽量补齐：project、project_en、client_type、industry、country、products_used、area_or_quantity、problem_solved、result、key_highlight、usable_in_outreach；官网没有明确数据的字段留空或写入 metadata.notes，不能脑补。
- customer_matching_guide 要具体到客户类型和开发信重点，不要写泛泛的“适合各类客户”。
- boundaries 必须明确 claims_we_can_make、claims_we_cannot_make、sensitive_topics，避免后续开发信夸大或乱写。
- metadata.profile_completeness 必须按字段完整度评估。只要已经识别出产品、案例、优势、客户类型等实质内容，不得返回 0。
{image_block}
=== 用户原始输入 ===
{user_message or source_text or "未提供"}

=== 公司官网 ===
{website or "未提供"}

=== 官网抓取资料 ===
{_compact_scraped_markdown(scraped_markdown) if scraped_markdown else "未抓取到官网资料"}

请严格返回 JSON。"""


IMAGE_DESC_PROMPT = """请仔细观察这张图片，提取以下信息（如果图片中包含的话）：
1. 公司名称/品牌名/Logo 文字
2. 产品名称和描述
3. 核心优势、资质认证
4. 案例信息（客户名、项目名、地区、规模）
5. 合作模式、联系方式
6. 行业类别
请用中文简要描述你看到的所有内容，不要编造图片中不存在的。"""


def _decode_uploaded_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return ""


def _extract_docx_text(raw_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""

    doc = Document(io.BytesIO(raw_bytes))
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_xlsx_text(raw_bytes: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ""

    workbook = load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets[:5]:
        lines.append(f"[Sheet: {sheet.title}]")
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if row_index > 200:
                lines.append("[已截断：仅读取前 200 行]")
                break
            values = [str(value).strip() for value in row if value not in (None, "")]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def _extract_uploaded_files_text(files: list[dict] | None) -> str:
    if not files:
        return ""

    parts: list[str] = []
    for file in files:
        filename = str(file.get("filename") or "uploaded-file")
        data = str(file.get("data") or "")
        if not data:
            continue

        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        try:
            raw_bytes = base64.b64decode(data)
        except Exception as exc:
            parts.append(f"### 文件：{filename}\n[无法解码上传文件：{exc}]")
            continue

        text = ""
        if ext in {"txt", "md", "csv"}:
            text = _decode_uploaded_text(raw_bytes)
            if ext == "csv" and text:
                sample = []
                reader = csv.reader(io.StringIO(text))
                for index, row in enumerate(reader):
                    if index >= 200:
                        sample.append("[已截断：仅读取前 200 行]")
                        break
                    sample.append(" | ".join(cell.strip() for cell in row))
                text = "\n".join(sample)
        elif ext == "docx":
            text = _extract_docx_text(raw_bytes)
        elif ext in {"xlsx", "xlsm"}:
            text = _extract_xlsx_text(raw_bytes)
        else:
            text = f"[暂不支持直接解析 .{ext or 'unknown'} 文件，请用户补充关键文字内容]"

        if text:
            parts.append(f"### 文件：{filename}\n{text[:12000]}")

    return "\n\n".join(parts)[:30000]


def _append_uploaded_file_text(params: dict, source_text: str) -> str:
    file_text = _extract_uploaded_files_text(params.get("files") or [])
    if not file_text:
        return source_text

    combined = "\n\n".join(
        part for part in [source_text, "=== 用户上传文件提取内容 ===\n" + file_text] if part
    )
    params["source_text"] = combined
    return combined


def _describe_images_sync(images_b64: list[str]) -> str:
    """Use vision model to describe uploaded images for profile building.

    Processes images one at a time to avoid write-timeout with large payloads.
    """
    if not images_b64:
        return ""
    client = replicate.Client(
        api_token=settings.replicate_api_token,
        timeout=httpx.Timeout(connect=30.0, read=120.0, write=300.0, pool=20.0),
    )
    descriptions: list[str] = []
    for idx, b64 in enumerate(images_b64):
        content_parts: list[dict] = [
            {"type": "text", "text": IMAGE_DESC_PROMPT},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
        ]
        try:
            output = client.run(
                settings.replicate_model,
                input={"messages": [{"role": "user", "content": content_parts}]},
            )
            desc = "".join(output) if isinstance(output, list) else str(output)
            descriptions.append(f"[图片 {idx + 1}]\n{desc}")
            logger.info("Image %d described successfully (%d chars)", idx + 1, len(desc))
        except Exception as exc:
            logger.warning("Image %d description failed: %s", idx + 1, exc)
    return "\n\n".join(descriptions)


def _analyze_profile_sync(params: dict, source_text: str, scraped_markdown: str, image_descriptions: str = "") -> dict:
    user_prompt = _build_user_prompt(params, source_text, scraped_markdown, image_descriptions)
    # Use shorter edit-focused system prompt when modifying existing profile
    is_edit = isinstance(params.get("existing_profile"), dict) and bool(params["existing_profile"])
    system_prompt = PROFILE_EDIT_SYSTEM_PROMPT if is_edit else PROFILE_SYSTEM_PROMPT
    client = replicate.Client(
        api_token=settings.replicate_api_token,
        timeout=httpx.Timeout(connect=20.0, read=420.0, write=60.0, pool=20.0),
    )
    last_error: Exception | None = None
    for attempt in range(3):
        try:
            output = client.run(
                settings.replicate_model,
                input={
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                },
            )
            response_text = "".join(output) if isinstance(output, list) else str(output)
            logger.info("AI response (%d chars): %s", len(response_text), response_text[:500])
            parsed = _parse_ai_json(response_text)
            if parsed is None:
                logger.warning("AI response not valid JSON (first 500 chars): %s", response_text[:500])
                if attempt < 2:
                    time.sleep(2)
                    continue
                raise ValueError(f"AI response is not valid JSON: {response_text[:200]}")
            return parsed
        except (httpx.TimeoutException, TimeoutError) as exc:
            last_error = exc
            if attempt < 2:
                time.sleep(3 * (attempt + 1))
                continue
            break

    raise RuntimeError(f"AI 整理画像超时，请稍后重试：{last_error}")


async def run_profile_pipeline(task_id: int, user_id: int, intent: dict) -> None:
    """Run company-profile collection as a background task."""
    params = intent.get("params", {})
    source_text = params.get("source_text", "")
    source_text = _append_uploaded_file_text(params, source_text)
    website = _normalize_url(params.get("url") or extract_url(source_text))
    if website:
        params["url"] = website

    try:
        async with async_session_factory() as db:
            await _create_task_log(
                db,
                task_id,
                step=1,
                name="分析画像需求",
                status="completed",
                message="已读取用户提供的公司信息和官网线索",
                progress=100,
            )
            task_manager.update_heartbeat(task_id)

            # Process uploaded images with vision model
            images_b64 = params.get("images") or []
            image_descriptions = ""
            if images_b64:
                try:
                    image_descriptions = await _run_blocking_step_with_progress(
                        db=db,
                        task_id=task_id,
                        step=1,
                        func=_describe_images_sync,
                        args=(images_b64,),
                        timeout_seconds=120,
                        running_message=f"AI 正在识别 {len(images_b64)} 张上传图片",
                        timeout_message="图片识别超时，将跳过图片继续整理",
                    )
                    if image_descriptions:
                        note = f"已识别 {len(images_b64)} 张图片内容"
                        await _update_task_log(
                            db, task_id, step=1,
                            message=f"已读取用户信息、官网线索{note}",
                        )
                except Exception as exc:
                    logger.warning("Image description failed: %s", exc)

            scraped_markdown = ""
            source_urls: list[str] = []
            await _create_task_log(
                db,
                task_id,
                step=2,
                name="抓取官网资料",
                status="running",
                message=f"正在抓取官网：{website}" if website else "未提供官网，跳过官网抓取",
                progress=0,
            )
            if website:
                try:
                    scraped_markdown, source_urls = await _run_blocking_step_with_progress(
                        db=db,
                        task_id=task_id,
                        step=2,
                        func=_scrape_website,
                        args=(website,),
                        timeout_seconds=180,
                        running_message=f"正在抓取官网：{website}",
                        timeout_message="官网抓取超过 3 分钟，已切换为使用用户输入继续整理",
                    )
                    await _update_task_log(
                        db,
                        task_id,
                        step=2,
                        status="completed",
                        message=f"已抓取 {len(source_urls)} 个官网页面，并优先深挖案例/项目页",
                        progress=100,
                    )
                except Exception as exc:
                    logger.warning("Company profile scrape failed: %s", exc)
                    await _update_task_log(
                        db,
                        task_id,
                        step=2,
                        status="completed",
                        message="官网抓取失败，改用用户输入继续整理",
                        progress=100,
                    )
                    user_material = (source_text or "").replace(website, "").strip()
                    user_material = re.sub(
                        r"https?://[A-Za-z0-9\-._~:/?#\[\]@!$&'()*+,;=%]+",
                        "",
                        user_material,
                    ).strip(" ，。,.!！?？：:")
                    if not user_material and not image_descriptions:
                        raise RuntimeError(
                            "官网抓取失败，且没有可用于整理画像的补充资料。"
                            "请确认网址可从服务器访问，或补充公司介绍、产品、优势、案例等文字资料。"
                        )
            else:
                await _update_task_log(
                    db,
                    task_id,
                    step=2,
                    status="completed",
                    message="未提供官网，已跳过",
                    progress=100,
                )
            task_manager.update_heartbeat(task_id)

            await _create_task_log(
                db,
                task_id,
                step=3,
                name="AI 整理画像",
                status="running",
                message="正在生成中文展示版和英文开发信素材版",
                progress=20,
            )
            try:
                profile = await _run_blocking_step_with_progress(
                    db=db,
                    task_id=task_id,
                    step=3,
                    func=_analyze_profile_sync,
                    args=(params, source_text, scraped_markdown, image_descriptions),
                    timeout_seconds=600,
                    running_message="AI 正在整理公司画像",
                    timeout_message="AI 整理画像超过 10 分钟，请稍后重试",
                )
            except Exception as exc:
                error_text = str(exc).lower()
                if "timeout" not in error_text and "timed out" not in error_text and "超时" not in str(exc):
                    raise
                logger.warning("AI profile analysis timed out; saving fallback draft: %s", exc)
                profile = _build_timeout_fallback_profile(
                    params,
                    source_text,
                    scraped_markdown,
                    source_urls,
                )
                await _update_task_log(
                    db,
                    task_id,
                    step=3,
                    status="completed",
                    message="AI 整理超时，已保存基础画像草稿，可补充资料后重新采集覆盖",
                    progress=100,
                )
            profile = _normalize_profile(profile, website, source_urls)
            markdown = _profile_to_markdown(profile)
            completeness = _normalize_completeness(
                profile.get("metadata", {}).get("profile_completeness")
            )
            await _update_task_log(
                db,
                task_id,
                step=3,
                status="completed",
                message=f"画像整理完成，完整度 {completeness:.0%}",
                progress=100,
            )
            task_manager.update_heartbeat(task_id)

            await _create_task_log(
                db,
                task_id,
                step=4,
                name="保存画像",
                status="running",
                message="正在保存为当前公司画像",
                progress=30,
            )
            task = await db.get(Task, task_id)
            profile_mode = params.get("profile_mode", "create")
            existing_profile_id = params.get("existing_profile_id")
            existing_profile = params.get("existing_profile")

            if profile_mode == "update" and existing_profile_id:
                # ── Update mode: merge into existing row ──
                existing_profile_dict = existing_profile if isinstance(existing_profile, dict) else {}
                merged_profile = _merge_profile_data(existing_profile_dict, profile)
                merged_profile = _normalize_profile(merged_profile, website, source_urls)
                merged_markdown = _profile_to_markdown(merged_profile)
                merged_completeness = _normalize_completeness(
                    merged_profile.get("metadata", {}).get("profile_completeness")
                )

                cp = await db.get(CompanyProfile, existing_profile_id)
                if cp:
                    cp.company_name = merged_profile.get("company_name", "")
                    cp.profile_data = merged_profile
                    cp.profile_markdown = merged_markdown
                    cp.task_id = task_id

                task.status = "completed"
                task.ended_at = datetime.now()
                task.result_summary = {
                    "companyName": merged_profile.get("company_name", ""),
                    "industry": merged_profile.get("industry", ""),
                    "products": len(merged_profile.get("products", [])),
                    "cases": len(merged_profile.get("case_studies", [])),
                    "completeness": merged_completeness,
                    "profileId": existing_profile_id,
                    "profileMode": "update",
                }
                await db.commit()
                await ensure_recommended_email_settings(
                    db, user_id, existing_profile_id, merged_profile
                )

                await _update_task_log(
                    db,
                    task_id,
                    step=4,
                    status="completed",
                    message="公司画像已更新（增量补充）",
                    progress=100,
                )
            else:
                # ── Create / Replace mode: new row ──
                await db.execute(
                    update(CompanyProfile)
                    .where(
                        CompanyProfile.user_id == user_id,
                        CompanyProfile.is_current == True,  # noqa: E712
                    )
                    .values(is_current=False)
                )
                await sync_company_profiles_id_sequence(db)
                cp = CompanyProfile(
                    user_id=user_id,
                    task_id=task_id,
                    company_name=profile.get("company_name", ""),
                    profile_data=profile,
                    profile_markdown=markdown,
                    is_current=True,
                )
                db.add(cp)

                task.status = "completed"
                task.ended_at = datetime.now()
                task.result_summary = {
                    "companyName": profile.get("company_name", ""),
                    "industry": profile.get("industry", ""),
                    "products": len(profile.get("products", [])),
                    "cases": len(profile.get("case_studies", [])),
                    "completeness": completeness,
                    "profileId": None,
                    "profileMode": "create",
                }
                await db.commit()
                await db.refresh(cp)
                await ensure_recommended_email_settings(db, user_id, cp.id, profile)
                task.result_summary = {
                    **(task.result_summary or {}),
                    "profileId": cp.id,
                }
                await db.commit()

                await _update_task_log(
                    db,
                    task_id,
                    step=4,
                    status="completed",
                    message="公司画像已保存",
                    progress=100,
                )
            task_manager.update_heartbeat(task_id)

    except asyncio.CancelledError:
        logger.info("Company profile pipeline cancelled for task %d", task_id)
        raise
    except Exception as exc:
        logger.exception("Company profile pipeline failed: %s", exc)
        async with async_session_factory() as db:
            task = await db.get(Task, task_id)
            if task:
                task.status = "failed"
                task.ended_at = datetime.now()
                error_text = str(exc)
                if "timed out" in error_text.lower() or "超时" in error_text:
                    error_text = "AI 整理画像超时，请稍后重试；如果官网内容较多，可以先提供公司简介、产品和案例的核心文字。"
                task.result_summary = {"error": error_text[:500]}
                await db.commit()
    finally:
        task_manager.remove_task(task_id)


async def run_supplement_pipeline(task_id: int, user_id: int, intent: dict) -> None:
    """Lightweight 3-step pipeline for supplementing an existing company profile.

    Step 1: Extract supplementary materials (image recognition + optional new URL scraping)
    Step 2: AI incremental supplement (existing_profile + new materials → complete profile JSON)
    Step 3: Save profile (directly overwrite existing row)
    """
    params = intent.get("params", {})
    source_text = params.get("source_text", "")
    source_text = _append_uploaded_file_text(params, source_text)
    existing_profile_id = params.get("existing_profile_id")

    if not existing_profile_id:
        logger.error("Supplement pipeline called without existing_profile_id for task %d", task_id)
        async with async_session_factory() as db:
            task = await db.get(Task, task_id)
            if task:
                task.status = "failed"
                task.ended_at = datetime.now()
                task.result_summary = {"error": "缺少已有画像ID，无法执行补充"}
                await db.commit()
        return

    existing_profile = params.get("existing_profile") or {}
    existing_website = existing_profile.get("website", "") if isinstance(existing_profile, dict) else ""

    try:
        async with async_session_factory() as db:
            # ── Step 1: Extract supplementary materials ──
            await _create_task_log(
                db, task_id, step=1,
                name="提取补充材料",
                status="running",
                message="正在提取补充材料...",
                progress=10,
            )
            task_manager.update_heartbeat(task_id)

            # 1a: Image recognition
            images_b64 = params.get("images") or []
            image_descriptions = ""
            if images_b64:
                try:
                    image_descriptions = await _run_blocking_step_with_progress(
                        db=db, task_id=task_id, step=1,
                        func=_describe_images_sync,
                        args=(images_b64,),
                        timeout_seconds=120,
                        running_message=f"AI 正在识别 {len(images_b64)} 张上传图片",
                        timeout_message="图片识别超时，将跳过图片继续补充",
                    )
                except Exception as exc:
                    logger.warning("Image description failed in supplement: %s", exc)

            # 1b: Optional new URL scraping (skip if same as existing website)
            new_url = _normalize_url(params.get("url") or extract_url(source_text))
            scraped_markdown = ""
            source_urls: list[str] = []
            if new_url and new_url != existing_website:
                try:
                    scraped_markdown, source_urls = await _run_blocking_step_with_progress(
                        db=db, task_id=task_id, step=1,
                        func=_scrape_website,
                        args=(new_url,),
                        timeout_seconds=180,
                        running_message=f"正在抓取新页面：{new_url}",
                        timeout_message="页面抓取超时，将使用已有材料继续补充",
                    )
                except Exception as exc:
                    logger.warning("Supplement scrape failed: %s", exc)

            parts = []
            if image_descriptions:
                parts.append(f"已识别 {len(images_b64)} 张图片")
            if scraped_markdown:
                parts.append(f"已抓取 {len(source_urls)} 个页面")
            if not parts:
                parts.append("提取用户输入文字")
            await _update_task_log(
                db, task_id, step=1,
                status="completed",
                message="、".join(parts),
                progress=100,
            )
            task_manager.update_heartbeat(task_id)

            # ── Step 2: AI incremental supplement ──
            await _create_task_log(
                db, task_id, step=2,
                name="AI 增量补充",
                status="running",
                message="AI 正在基于现有画像增量补充新资料",
                progress=20,
            )
            try:
                profile = await _run_blocking_step_with_progress(
                    db=db, task_id=task_id, step=2,
                    func=_analyze_profile_sync,
                    args=(params, source_text, scraped_markdown, image_descriptions),
                    timeout_seconds=600,
                    running_message="AI 正在增量补充公司画像",
                    timeout_message="AI 补充画像超过 10 分钟，请稍后重试",
                )
                await _update_task_log(
                    db, task_id, step=2,
                    status="completed",
                    message="画像补充完成",
                    progress=100,
                )
            except Exception as exc:
                error_text = str(exc).lower()
                if "timeout" not in error_text and "timed out" not in error_text and "超时" not in str(exc):
                    raise
                logger.warning("AI supplement timed out for task %d, keeping existing profile: %s", task_id, exc)
                async with async_session_factory() as err_db:
                    task = await err_db.get(Task, task_id)
                    if task:
                        task.status = "failed"
                        task.ended_at = datetime.now()
                        task.result_summary = {"error": "AI 补充超时，现有画像已保留不变"}
                        await err_db.commit()
                return

            task_manager.update_heartbeat(task_id)

            # ── Step 3: Save profile (overwrite existing row) ──
            await _create_task_log(
                db, task_id, step=3,
                name="保存画像",
                status="running",
                message="正在保存更新后的公司画像",
                progress=30,
            )

            # Merge AI incremental result into existing profile
            if isinstance(existing_profile, dict) and existing_profile:
                ai_keys = set(profile.keys())
                full_keys = {"company_name", "one_line_intro", "products", "certifications",
                             "case_studies", "core_competencies", "english_profile", "metadata"}
                if not full_keys.issubset(ai_keys):
                    logger.info("Supplement AI returned partial JSON (%d keys), merging into existing profile", len(profile))
                profile = _merge_profile_data(existing_profile, profile)

            profile = _normalize_profile(profile, existing_website, source_urls)
            completeness = _normalize_completeness(
                profile.get("metadata", {}).get("profile_completeness")
            )
            markdown = _profile_to_markdown(profile)
            cp = await db.get(CompanyProfile, existing_profile_id)
            if cp:
                cp.company_name = profile.get("company_name", "")
                cp.profile_data = profile
                cp.profile_markdown = markdown
                cp.task_id = task_id

            task = await db.get(Task, task_id)
            task.status = "completed"
            task.ended_at = datetime.now()
            task.result_summary = {
                "companyName": profile.get("company_name", ""),
                "industry": profile.get("industry", ""),
                "products": len(profile.get("products", [])),
                "cases": len(profile.get("case_studies", [])),
                "completeness": completeness,
                "profileId": existing_profile_id,
                "profileMode": "update",
            }
            await db.commit()
            await ensure_recommended_email_settings(db, user_id, existing_profile_id, profile)

            await _update_task_log(
                db, task_id, step=3,
                status="completed",
                message=f"公司画像已更新（增量补充），完整度 {completeness:.0%}",
                progress=100,
            )
            task_manager.update_heartbeat(task_id)

    except asyncio.CancelledError:
        logger.info("Supplement pipeline cancelled for task %d", task_id)
        raise
    except Exception as exc:
        logger.exception("Supplement pipeline failed for task %d: %s", task_id, exc)
        async with async_session_factory() as db:
            task = await db.get(Task, task_id)
            if task:
                task.status = "failed"
                task.ended_at = datetime.now()
                task.result_summary = {"error": str(exc)[:500]}
                await db.commit()
    finally:
        task_manager.remove_task(task_id)


async def run_profile_quick_edit(task_id: int, user_id: int, intent: dict) -> None:
    """Single-step pipeline for quick profile edits (text/images, no URL scraping).

    Used when the user only provides text, images, or files to modify an existing
    profile — no website scraping needed. Runs in one step: describe images (if
    any) → AI call with existing profile → save.
    """
    params = intent.get("params", {})
    source_text = params.get("source_text", "")
    source_text = _append_uploaded_file_text(params, source_text)
    existing_profile_id = params.get("existing_profile_id")

    if not existing_profile_id:
        logger.error("Quick edit called without existing_profile_id for task %d", task_id)
        async with async_session_factory() as db:
            task = await db.get(Task, task_id)
            if task:
                task.status = "failed"
                task.ended_at = datetime.now()
                task.result_summary = {"error": "缺少已有画像ID，无法执行修改"}
                await db.commit()
        return

    try:
        async with async_session_factory() as db:
            await _create_task_log(
                db, task_id, step=1,
                name="修改画像",
                status="running",
                message="正在根据您的修改意见更新画像...",
                progress=10,
            )
            task_manager.update_heartbeat(task_id)

            # Describe images if any (runs within the same step)
            images_b64 = params.get("images") or []
            image_descriptions = ""
            if images_b64:
                try:
                    image_descriptions = await _run_blocking_step_with_progress(
                        db=db, task_id=task_id, step=1,
                        func=_describe_images_sync,
                        args=(images_b64,),
                        timeout_seconds=120,
                        running_message=f"AI 正在识别 {len(images_b64)} 张上传图片",
                        timeout_message="图片识别超时，将跳过图片继续修改",
                    )
                except Exception as exc:
                    logger.warning("Image description failed in quick_edit: %s", exc)

            # AI call — existing_profile is already in params
            try:
                profile = await _run_blocking_step_with_progress(
                    db=db, task_id=task_id, step=1,
                    func=_analyze_profile_sync,
                    args=(params, source_text, "", image_descriptions),
                    timeout_seconds=600,
                    running_message="AI 正在修改画像",
                    timeout_message="AI 修改超时，现有画像已保留不变",
                )
            except Exception as exc:
                error_text = str(exc).lower()
                if "timeout" not in error_text and "timed out" not in error_text and "超时" not in str(exc):
                    raise
                logger.warning("AI quick edit timed out for task %d: %s", task_id, exc)
                async with async_session_factory() as err_db:
                    task = await err_db.get(Task, task_id)
                    if task:
                        task.status = "failed"
                        task.ended_at = datetime.now()
                        task.result_summary = {"error": "AI 修改超时，现有画像已保留不变"}
                        await err_db.commit()
                return

            existing_website = ""
            ep = params.get("existing_profile")
            if isinstance(ep, dict):
                existing_website = ep.get("website", "")

            # AI returns only modified fields — directly replace, NOT merge
            # (merge would undo deletions by adding old items back)
            if isinstance(ep, dict) and ep:
                ai_keys = set(profile.keys())
                full_keys = {"company_name", "one_line_intro", "products", "certifications",
                             "case_studies", "core_competencies", "english_profile", "metadata"}
                if not full_keys.issubset(ai_keys):
                    # Partial response — replace modified fields into existing profile
                    logger.info("AI returned partial JSON (%d keys): %s", len(profile), list(profile.keys()))
                    for key, value in profile.items():
                        ep[key] = value
                    profile = ep
            profile = _normalize_profile(profile, existing_website, [])
            completeness = _normalize_completeness(
                profile.get("metadata", {}).get("profile_completeness")
            )
            markdown = _profile_to_markdown(profile)

            # Save — overwrite existing row
            cp = await db.get(CompanyProfile, existing_profile_id)
            if cp:
                cp.company_name = profile.get("company_name", "")
                cp.profile_data = profile
                cp.profile_markdown = markdown
                cp.task_id = task_id

            task = await db.get(Task, task_id)
            task.status = "completed"
            task.ended_at = datetime.now()
            task.result_summary = {
                "companyName": profile.get("company_name", ""),
                "industry": profile.get("industry", ""),
                "products": len(profile.get("products", [])),
                "cases": len(profile.get("case_studies", [])),
                "completeness": completeness,
                "profileId": existing_profile_id,
                "profileMode": "update",
            }
            await db.commit()
            await ensure_recommended_email_settings(db, user_id, existing_profile_id, profile)

            await _update_task_log(
                db, task_id, step=1,
                status="completed",
                message=f"画像修改完成，完整度 {completeness:.0%}",
                progress=100,
            )
            task_manager.update_heartbeat(task_id)

    except asyncio.CancelledError:
        logger.info("Quick edit cancelled for task %d", task_id)
        raise
    except Exception as exc:
        logger.exception("Quick edit failed for task %d: %s", task_id, exc)
        async with async_session_factory() as db:
            task = await db.get(Task, task_id)
            if task:
                task.status = "failed"
                task.ended_at = datetime.now()
                task.result_summary = {"error": str(exc)[:500]}
                await db.commit()
    finally:
        task_manager.remove_task(task_id)
